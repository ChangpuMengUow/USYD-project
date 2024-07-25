from docx import Document
import base64
import io

def getdoc(file_content_base64):
    # Decode the base64 content to binary
    file_content = base64.b64decode(file_content_base64)
    
    # Load the DOCX file from the binary content
    doc = Document(io.BytesIO(file_content))
    
    # Create a new DOCX document for the extracted content
    new_doc = Document()
    
    # Iterate over the tables in the document
    for table in doc.tables:
        # Find the index of the "Content" header
        content_col_index = None
        for i, cell in enumerate(table.rows[0].cells):
            if "Content" in cell.text:
                content_col_index = i
                break
        
        if content_col_index is not None:
            # Iterate over the rows in the table starting from the second row
            for row in table.rows[1:]:
                # Extract the content from the "Content" column
                content_text = row.cells[content_col_index].text
                # Add the content to the new document
                new_doc.add_paragraph(content_text)
    
    # Save the new document to a binary stream
    new_doc_stream = io.BytesIO()
    new_doc.save(new_doc_stream)
    new_doc_stream.seek(0)
    
    # Encode the new document to base64
    new_doc_base64 = base64.b64encode(new_doc_stream.read()).decode('utf-8')
    
    return new_doc_base64


def process_excel(file_content_base64):
    file_content = base64.b64decode(file_content_base64)
    excel = pd.read_excel(BytesIO(file_content))
    row_count = len(excel)
    return row_count

# Example usage
if __name__ == "__main__":
    # This is just for testing the function locally.
    with open("path_to_your_docx_file.docx", "rb") as file:
        file_content = file.read()
        file_content_base64 = base64.b64encode(file_content).decode('utf-8')
        result = getdoc(file_content_base64)
        with open("extracted_content.docx", "wb") as output_file:
            output_file.write(base64.b64decode(result))
